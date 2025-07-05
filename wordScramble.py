"""
The purpose of this program is to create a word scramble document using words pulled from existing lists in the SpellingQuiz repository (https://github.com/Monkiko/SpellingQuiz).

Created by: Ian Rivera-Leandry
Last Updated: July 4, 2025
Version: 0.1.0
"""

import os
import urllib.request
from random import shuffle
from time import sleep
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor


# Get input from the user to determine the grade level and the number of the list to pull from

def start():
    grade_level = str(input("Please indicate the grade level of the list being used (i.e. K, 1, 2, 3, etc.): "))
    list_number = str(input("Please indicate the list number of the list being used: "))

    if len(grade_level) == 1:
        readSpellingList(grade_level, list_number)

    else:
        clean()
        print("Incorrect input. Please enter either K, 1, 2, 3, etc.")
        sleep(3)
        start()


# Clear the screen dependent on the host OS system
def clean():
    if os.name == "nt":
        _ = os.system("cls")
    else:
        _ = os.system("clear")


def readSpellingList(grade_level, list_number):

    match grade_level.lower():
        case "k":
            grade = "Kindergarten"
        case "1":
            grade = "First_Grade"
        case "2":
            grade = "Second_Grade"
        case "3":
            grade = "Third_Grade"
        case "4":
            grade = "Fourth_Grade"
        case "5":
            grade = "Fifth_Grade"
        case "6":
            grade = "Sixth_Grade"
        case "7":
            grade = "Seventh_Grade"
        case "8":
            grade = "Eighth_Grade"
    
    target_url = "https://raw.githubusercontent.com/Monkiko/SpellingQuiz/refs/heads/main/Tests/" + grade + "/Simplified_Lists/Spelling_List_%23" + list_number + ".txt"
    print(target_url)
    with urllib.request.urlopen(target_url) as response:
        spellingList = response.read().decode().splitlines()

    shuffle(spellingList)
    scrambleWords(spellingList, grade, list_number)


def scrambleWords(spellingList, grade, list_number):
    
    scrambledList = []
    for word in spellingList:
        word = list(word)
        shuffle(word)
        word = ''.join(word)
        scrambledList.append(word)

    printScrambledList(scrambledList, grade, list_number)


def printScrambledList(scrambledList, grade, list_number):
    document = Document()

    title = document.add_heading("Spelling List #" + list_number + " Word Scramble", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.name = "Arial"
    title.style.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    title.style.font.size = Pt(26)

    for i in range(10):
        list = document.add_paragraph(str(i + 1) + ") " + scrambledList[i].strip() + "\t\t\t\t______")
        list.style.font.name = "Arial"
        list.style.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        list.style.font.size = Pt(20)

    document.save("Scrambled_Lists/" + grade + "/Spelling List #" + list_number + " WS.docx")
    print("Spelling list saved to Scrambled_Lists/" + grade + "/Spelling List #" + list_number + " WS.docx")

start()